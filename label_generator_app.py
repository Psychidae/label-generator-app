import streamlit as st
from streamlit_folium import st_folium
import folium
import requests
import pandas as pd
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re
import json

# --- Constants ---
# Region Color Mapping
REGION_COLORS = {
    "Palearctic (White)": "#FFFFFF",
    "Oriental (Yellow)": "#FFFF00",
    "Wallacea/Melanesia (Orange)": "#FFA500",
    "Australian (Red)": "#FF0000",
    "NZ/Pacific (Brown)": "#964B00",
    "Nearctic (Green)": "#008000",
    "Neotropical (Yellow-Green)": "#9ACD32",
    "Ethiopian (Blue)": "#0000FF",
    "Madagascar (Purple)": "#800080"
}

def generate_html_sheet(queue, num_columns, font_name, font_size, label_color):
    """Generates an HTML representation of the full A4 sheet."""
    
    # CSS for A4 Sheet and Grid
    # A4 is 210mm x 297mm.
    # Grid columns = num_columns.
    
    css = f"""
    <style>
        @page {{ size: A4; margin: 0; }}
        .sheet {{
            width: 210mm;
            min-height: 297mm;
            padding: 5mm; /* Margins */
            box-sizing: border-box;
            background: white;
            border: 1px solid #eee;
            margin: 0 auto;
            display: grid;
            grid-template-columns: repeat({num_columns}, 1fr);
            grid-auto-rows: min-content;
            font-family: "{font_name}", Arial, sans-serif;
        }}
        .cell {{
            border: 1px dotted #CCCCCC; /* Dotted Gray */
            padding: 1px;
            box-sizing: border-box;
            overflow: hidden;
            font-size: {font_size}pt;
            line-height: 1.1;
        }}
        .header {{ font-weight: bold; }}
        .bar {{ height: 2px; margin: 1px 0; }}
        .body {{ white-space: pre-wrap; }}
    </style>
    """
    
    # Build Cells
    cells_html = ""
    
    # Flatten items
    all_items = []
    for item in queue:
        for _ in range(item['quantity']):
            all_items.append(item)
            
    for item in all_items:
        ctype = item.get('type', 'text')
        content_html = ""
        
        if ctype == 'data_v2':
            # Use item color or default
            i_color = item.get('color', '#000000')
            content_html = f"""
                <div class="header">{item['header']}</div>
                <div class="bar" style="background-color: {i_color};"></div>
                <div class="body">{item['body']}</div>
            """
        elif ctype == 'rich':
            # Simplified rich text render for preview
            # (In a real full implementation, we'd parse the list of tuples)
            preview_txt = item['preview']
            content_html = f"<div>{preview_txt}</div>"
        else:
             content_html = f"<div>{str(item.get('content', ''))}</div>"
             
        cells_html += f'<div class="cell">{content_html}</div>'

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>{css}</head>
    <body>
        <div class="sheet">
            {cells_html}
        </div>
    </body>
    </html>
    """
    return html

# Tabs
tab1, tab2, tab3, tab4 = st.tabs(["🌎 Data Label", "🔍 Identification Label", "🧬 Molecular Label", "📄 Sheet Preview"])

# ... (Tab 1, 2, 3 Logic remains same, only indentation/context might shift but since I'm appending tab4, let's see) ...
# Actually, I need to update the `st.tabs` call which is way up in the file. 
# And add the Tab 4 content at end of file (or appropriate place).

# Strategy: Replace the `tab1, tab2, tab3 = ...` line with 4 tabs.
# Then append Tab 4 logic at end of file.

# Step 1: Replace Tabs line (I'll do this in a separate chunk via context match if needed, or if I can assume location)
# Step 2: Add generate_html_sheet definition (Top of file or near helpers)
# Step 3: Add Tab 4 content.

# Let's do Step 1 & 3 here, assuming I can find the tabs line.
# Actually, I'll put generate_html_sheet near create_docx.



# --- Configuration ---
# API endpoints
GEOCODING_API_ENDPOINT = "https://maps.googleapis.com/maps/api/geocode/json"
ELEVATION_API_ENDPOINT = "https://maps.googleapis.com/maps/api/elevation/json"

# Default API Key (Securely loaded from secrets)
# When running locally, create .streamlit/secrets.toml
# When running on Streamlit Cloud, set this in the App Settings
try:
    DEFAULT_API_KEY = st.secrets["GOOGLE_MAPS_API_KEY"]
except FileNotFoundError:
    DEFAULT_API_KEY = "" # Fallback if no secrets file found
except KeyError:
    DEFAULT_API_KEY = "" # Fallback if key missing


# --- Helper Functions (Adapted from label_app.py) ---
def get_elevation(lat, lon, api_key):
    """Calls the Google Elevation API to get the altitude. Returns None if invalid."""
    if not api_key: return None
    params = {'locations': f'{lat},{lon}', 'key': api_key}
    try:
        response = requests.get(ELEVATION_API_ENDPOINT, params=params, timeout=10)
        data = response.json()
        if data['status'] == 'OK' and len(data['results']) > 0:
            return int(round(data['results'][0]['elevation']))
    except:
        pass
    return None

def to_roman(n):
    """Converts 1-12 to Roman numerals."""
    return ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII"][n]

def format_coordinates_v2(lat, lon):
    """Formats coordinates as 35.689°N, 139.691°E"""
    lat_dir = "N" if lat >= 0 else "S"
    lon_dir = "E" if lon >= 0 else "W"
    return f"{abs(lat):.3f}°{lat_dir}, {abs(lon):.3f}°{lon_dir}"

def get_google_address_struct(lat, lon, api_key):
    """
    Returns a dict with structured address components.
    """
    if not api_key: return None
    params = {'latlng': f'{lat},{lon}', 'key': api_key, 'language': 'en'}
    try:
        response = requests.get(GEOCODING_API_ENDPOINT, params=params, timeout=10)
        data = response.json()
    except:
        return None
        
    if data['status'] == 'OK' and len(data['results']) > 0:
        result = data['results'][0]
        components = result.get('address_components', [])
        
        addr_info = {'country': '', 'admin': '', 'locality': ''}
        
        for c in components:
            types = c.get('types', [])
            if 'country' in types:
                addr_info['country'] = c.get('short_name', '').upper() # Country Code or Name? User requested CAPS. Short name is usually code (JP). Long name is JAPAN. User style showed MADAGASCAR. Let's use long_name and Upper it.
                if 'long_name' in c: addr_info['country'] = c.get('long_name', '').upper()
            if 'administrative_area_level_1' in types:
                addr_info['admin'] = c.get('long_name', '')
            if 'locality' in types:
                addr_info['locality'] = c.get('long_name', '')
            if not addr_info['locality'] and 'administrative_area_level_2' in types:
                 addr_info['locality'] = c.get('long_name', '')
                 
        return addr_info
    return None

def generate_label_body_v2(locality, elev, lat, lon, date_obj, collector, method):
    """Generates the body text (excluding header) for the new style."""
    
    # Line 1 (Body): Locality, (alt. XXX m),
    parts_L1 = []
    if locality: parts_L1.append(locality)
    if elev is not None: parts_L1.append(f"(alt. {elev} m)")
    line1 = ", ".join(parts_L1)
    if line1 and not line1.endswith(','): line1 += ","

    # Line 2: Coords, Date
    coords = format_coordinates_v2(lat, lon)
    # Date: 7-12 II 2018. We have single date. -> 15 II 2023
    d_str = f"{date_obj.day} {to_roman(date_obj.month)} {date_obj.year}"
    line2 = f"{coords}, {d_str},"
    
    # Line 3: Collector, (Method)
    line3 = collector
    if method:
        # Check if method already hasparens
        if method.startswith('(') and method.endswith(')'):
            line3 += f", {method}"
        else:
            line3 += f", ({method})"
            
    return f"{line1}\n{line2}\n{line3}"

# Keep old helpers for backward compatibility if needed, or replace. 
# For safety, I'll modify existing calls in the main loop to use these new V2 functions.

# --- Helper Functions (New) ---

def parse_coordinates(coord_string):
    """
    Parses a string to extract latitude and longitude.
    Supports formats like:
    - "35.123, 139.456"
    - "35.123 139.456"
    - "N35.123 E139.456"
    """
    coord_string = coord_string.replace(',', ' ').replace('N', '').replace('E', '').replace('n', '').replace('e', '')
    parts = coord_string.split()
    if len(parts) >= 2:
        try:
            return float(parts[0]), float(parts[1])
        except ValueError:
            return None, None
    return None, None

def set_paragraph_shading(paragraph, color_hex):
    """Sets the background shading of a paragraph."""
    val = color_hex.replace("#", "")
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), val)
    paragraph._p.get_or_add_pPr().append(shd)

def set_run_spacing(run, value_pt):
    """Sets character spacing (kerning/condensing). Value in points."""
    if value_pt == 0: return
    val = int(value_pt * 20)
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(val))
    rPr.append(spacing)

def set_table_properties(table, show_borders=True):
    """
    Sets custom table properties:
    1. Borders: Dotted Light Gray (#CCCCCC) if show_borders is True.
    2. Margins: 0 for max density.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # 1. Borders
    if show_borders:
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'dotted')
            border.set(qn('w:sz'), '4') # 1/8 pt, minimal vis
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CCCCCC') # Light Gray
            tblBorders.append(border)
        tblPr.append(tblBorders)
    
    # 2. Cell Margins (Zero)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        width = OxmlElement(f'w:{side}')
        width.set(qn('w:w'), '0')
        width.set(qn('w:type'), 'dxa')
        tblCellMar.append(width)
    tblPr.append(tblCellMar)

def create_docx(label_queue, font_size=4.0, show_borders=True, num_columns=13, font_name='Arial', char_spacing=0.0):
    """
    Creates a DOCX file from a list of label objects using a Grid Layout (Table).
    Optimized for insect specimens (small font, efficient cutting).
    """
    doc = Document()
    
    # A4 Setup (Narrow margins to maximize printing area)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.2) 
    section.right_margin = Inches(0.2)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    
    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0 # Single spacing

    # Flatten queue into individual labels
    all_labels = []
    for item in label_queue:
        for _ in range(item['quantity']):
            all_labels.append(item) 

    # Create Table
    COLS = num_columns
    rows = -(-len(all_labels) // COLS) # Ceiling division
    
    if rows > 0:
        table = doc.add_table(rows=rows, cols=COLS)
        set_table_properties(table, show_borders)
    else:
        return io.BytesIO()

    # Populate Cells
    for idx, item in enumerate(all_labels):
        row = idx // COLS
        col = idx % COLS
        cell = table.cell(row, col)
        
        # Access the first paragraph (default) or add one
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0) # Tighter packing
        
        ctype = item.get('type', 'text')
        
        if ctype == 'data_v2':
            # 1. Header (Bold)
            run_h = p.add_run(item['header'])
            run_h.font.name = font_name
            run_h.font.size = Pt(font_size)
            run_h.bold = True
            set_run_spacing(run_h, char_spacing)
            
            # 2. Colored Bar
            p_bar = cell.add_paragraph()
            p_bar.paragraph_format.space_after = Pt(0)
            p_bar.paragraph_format.line_spacing = Pt(2) 
            run_bar = p_bar.add_run(" " * 5)
            run_bar.font.size = Pt(1.5) 
            set_paragraph_shading(p_bar, item['color'])
            
            # 3. Body
            p_body = cell.add_paragraph()
            p_body.paragraph_format.space_after = Pt(0) # Zero spacing
            run_b = p_body.add_run(item['body'])
            run_b.font.name = font_name
            run_b.font.size = Pt(font_size)
            set_run_spacing(run_b, char_spacing)
            
        elif ctype == 'rich':
            content = item['content']
            for segment, is_italic in content:
                run = p.add_run(segment)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.italic = is_italic
                set_run_spacing(run, char_spacing)
        else:
            # content is simple string
            content = item['content'] if 'content' in item else item.get('text', '')
            run = p.add_run(str(content))
            run.font.name = font_name
            run.font.size = Pt(font_size)
            set_run_spacing(run, char_spacing)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Main App ---

st.set_page_config(page_title="Specimen Label Generator", layout="wide")

st.title("🏷️ Specimen Label Generator")

# Initialize Session State
if 'lat' not in st.session_state: st.session_state.lat = 0.0
if 'lon' not in st.session_state: st.session_state.lon = 0.0
if 'last_map_click' not in st.session_state: st.session_state.last_map_click = None
if 'label_queue' not in st.session_state: st.session_state.label_queue = []
if 'last_fetched_coords' not in st.session_state: st.session_state.last_fetched_coords = (None, None)
if 'address_input' not in st.session_state: st.session_state.address_input = ""
if 'elevation_val' not in st.session_state: st.session_state.elevation_val = None

# Sidebar for Settings
with st.sidebar:
    st.header("Settings")
    api_key = st.text_input("Google Maps API Key", value=DEFAULT_API_KEY, type="password")
    
    st.divider()
    st.subheader("Print Settings")
    quantity = st.number_input("Quantity", min_value=1, value=1, step=1)
    
    # Font & Grid Settings
    font_size = st.slider("Font Size (pt)", min_value=3.0, max_value=8.0, value=4.0, step=0.5)
    font_name = st.selectbox("Font Family", ["Arial", "PT Sans Narrow", "Seravek", "Hiragino Sans", "Times New Roman"])
    char_spacing = st.slider("Character Spacing (pt)", min_value=-1.5, max_value=1.5, value=-0.5, step=0.1, help="Negative values condense text.")
    
    show_borders = st.checkbox("Show Grid Borders", value=True, help="Useful for cutting labels.")
    num_columns = st.number_input("Columns per Row", min_value=1, max_value=20, value=13, help="Adjust for width (13 cols approx 14mm width)")
    
    st.divider()
    st.subheader("🎨 Style (Region/Color)")
    
    # Region Selection
    region_options = list(REGION_COLORS.keys()) + ["Custom"]
    selected_region = st.selectbox("Select Biogeographic Region", region_options)
    
    if selected_region == "Custom":
        label_color = st.color_picker("Custom Bar Color", value="#6A0DAD")
    else:
        label_color = REGION_COLORS[selected_region]
        st.color_picker("Color Preview", value=label_color, disabled=True)
    
    # Batch Update Button (To fix the 'Printscreen setting not applied' issue for colors)
    if st.button("Apply Color to All Queued Items"):
        count = 0
        for item in st.session_state.label_queue:
            if item.get('type') == 'data_v2':
                item['color'] = label_color
                count += 1
        st.success(f"Updated color for {count} items.")
        st.rerun()

    
    st.divider()
    st.info("Paste Coordinates Example:\n35.689, 139.691")

    
    st.divider()
    st.divider()
    st.subheader("📦 Label Queue")
    
    # Save/Load System
    col_sl1, col_sl2 = st.columns(2)
    with col_sl1:
        # Save
        if st.session_state.label_queue:
            queue_json = json.dumps(st.session_state.label_queue, indent=2)
            st.download_button(
                label="💾 Save Queue Data (JSON)",
                data=queue_json,
                file_name=f"label_data_backup_{datetime.date.today()}.json",
                mime="application/json",
            )
            
    with col_sl2:
        # Load
        uploaded_file = st.file_uploader("📂 Load Queue Data", type=["json"])
        if uploaded_file is not None:
             try:
                 loaded_data = json.load(uploaded_file)
                 if isinstance(loaded_data, list):
                     if st.button("Confirm Load", type="primary"):
                         st.session_state.label_queue = loaded_data
                         st.success("Data Loaded!")
                         st.rerun()
                 else:
                     st.error("Invalid JSON format (must be a list).")
             except Exception as e:
                 st.error(f"Error loading JSON: {e}")

    if st.session_state.label_queue:
        st.write(f"Items in queue: {len(st.session_state.label_queue)}")
        if st.button("Clear Queue", type="secondary"):
            st.session_state.label_queue = []
            st.rerun()
    else:
        st.write("Queue is empty.")

# Tabs for Modules
tab1, tab2, tab3, tab4 = st.tabs(["🌎 Data Label", "🔍 Identification Label", "🧬 Molecular Label", "📄 Sheet Preview"])

# Initialize extra session states for V2 if not present
if 'header_input' not in st.session_state: st.session_state.header_input = ""
if 'locality_input' not in st.session_state: st.session_state.locality_input = ""

# --- TAB 1: DATA LABEL (Existing Internal Logic) ---
with tab1:
    col1, col2 = st.columns([1.5, 1])

    with col1:
        st.subheader("📍 Location")
        
        # Callback for Paste
        def on_paste_change():
            val = st.session_state.paste_coords
            if val:
                p_lat, p_lon = parse_coordinates(val)
                if p_lat is not None:
                    st.session_state.lat = p_lat
                    st.session_state.lon = p_lon
                    st.toast(f"Coordinates Updated: {p_lat}, {p_lon}")
                else:
                    st.toast("Invalid coordinate format", icon="⚠️")

        # Coordinate Paste Input
        st.text_input("Paste Coordinates (Lat, Lon)", key="paste_coords", placeholder="e.g. 35.6586, 139.7454", on_change=on_paste_change)
        
        # Map
        m = folium.Map(location=[36.2048, 138.2529], zoom_start=5)
        output = st_folium(m, height=400, use_container_width=True)

        # Logic to update state from Map Click
        if output and output['last_clicked'] != st.session_state.last_map_click:
            st.session_state.last_map_click = output['last_clicked']
            if output['last_clicked']:
                st.session_state.lat = output['last_clicked']['lat']
                st.session_state.lon = output['last_clicked']['lng']
                st.toast(f"Coordinates Updated from Map")

    with col2:
        st.subheader("📝 Details")
        
        # Coordinates Display/Edit (Bound to Session State)
        st.number_input("Latitude", format="%.6f", key="lat") 
        st.number_input("Longitude", format="%.6f", key="lon") 
        
        # --- Auto-Fetch Logic V2 ---
        current_coords = (st.session_state.lat, st.session_state.lon)
        if current_coords != st.session_state.last_fetched_coords:
            if api_key and not (current_coords[0] == 0.0 and current_coords[1] == 0.0):
                with st.spinner("Fetching Info..."):
                    addr_struct = get_google_address_struct(current_coords[0], current_coords[1], api_key)
                    elev = get_elevation(current_coords[0], current_coords[1], api_key)
                    
                    if addr_struct:
                        # Construct Header: COUNTRY: Region,
                        parts = []
                        if addr_struct['country']: parts.append(addr_struct['country'])
                        header_str = f"{addr_struct['country']}: {addr_struct['admin']},"
                        locality_str = addr_struct['locality']
                    else:
                        header_str = "COUNTRY: Region,"
                        locality_str = "Locality Not Found"

                    st.session_state.header_input = header_str
                    st.session_state.locality_input = locality_str
                    # Update manual elevation field
                    st.session_state.elevation_manual = str(elev) if elev is not None else ""
                    st.session_state.last_fetched_coords = current_coords
        
        # Inputs (V2 Fields)
        st.text_input("Header (Bold)", key="header_input", help="Format: COUNTRY: Region,")
        st.text_area("Locality", key="locality_input", height=60)
        
        # Elevation Input (Manual Override)
        if 'elevation_manual' not in st.session_state:
             st.session_state.elevation_manual = ""

        st.text_input("Elevation (m)", key="elevation_manual", help="Auto-filled if available, or enter manually.")
        
        collection_date = st.date_input("Collection Date", datetime.date.today())
        collector_name = st.text_input("Collector Name", value="M. Tsuchioka") 
        options = ["", "Light trap", "Sweeping", "Beating", "Bait trap", "Hand picking", "Fit", "Malaise trap"]
        collection_method = st.selectbox("Collection Method", options + ["Other"])
        if collection_method == "Other":
            collection_method = st.text_input("Enter Method")
            
        col_btn1, col_btn2 = st.columns(2)
        preview_btn = col_btn1.button("Preview Data Label", type="primary", use_container_width=True)
        add_queue_btn = col_btn2.button("Add Data Label", use_container_width=True)

    # Logic for Tab 1
    current_lat = st.session_state.lat
    current_lon = st.session_state.lon
    
    if preview_btn or add_queue_btn:
        if current_lat == 0.0 and current_lon == 0.0:
            st.warning("Please define a location.")
        else:
            final_header = st.session_state.header_input
            final_locality = st.session_state.locality_input
            final_elevation = st.session_state.elevation_manual
            
            # Generate Body Text
            body_text = generate_label_body_v2(
                final_locality, final_elevation, current_lat, current_lon,
                collection_date, collector_name, collection_method
            )
            
            # Preview HTML for V2
            preview_html = f"""
            <div style="
                border: 1px solid #000; padding: 5px; width: 250px; 
                font-family: Arial; font-size: 11px; line-height: 1.2;
                background: white; color: black;
            ">
                <div style="font-weight: bold;">{final_header}</div>
                <div style="height: 4px; background-color: {label_color}; margin: 2px 0;"></div>
                <div style="white-space: pre-wrap;">{body_text}</div>
            </div>
            """

            if preview_btn:
                st.info("Preview:")
                st.components.v1.html(preview_html, height=150)

            if add_queue_btn:
                st.session_state.label_queue.append({
                    'type': 'data_v2',
                    'header': final_header,
                    'body': body_text,
                    'color': label_color,
                    'quantity': quantity,
                    'preview': f"{final_header} {final_locality}..."
                })
                st.success(f"Added {quantity} Data Label(s) to Queue!")

# --- TAB 2: IDENTIFICATION LABEL ---
with tab2:
    st.header("Identification Label")
    col_id1, col_id2 = st.columns(2)
    
    with col_id1:
        family = st.text_input("Family")
        genus = st.text_input("Genus")
        species = st.text_input("Species")
        author = st.text_input("Author")
    
    with col_id2:
        det_name = st.text_input("Determined By (Name)", value="M. Tsuchioka")
        det_year = st.text_input("Determined Year", value=str(datetime.date.today().year))
        
    add_id_btn = st.button("Add Identification Label", use_container_width=True)
    
    if add_id_btn:
        # Construct Rich Content for DOCX: List of (text, italic)
        # Format:
        # Family (if present)
        # Genus species Author
        # det. Name Year
        
        rich_content = []
        preview_str = ""
        
        if family:
            rich_content.append((f"{family}\n", False))
            preview_str += f"{family}\n"
            
        if genus:
            rich_content.append((f"{genus} ", True)) # Genus Italic
            preview_str += f"*{genus}* "
        
        if species:
            rich_content.append((f"{species} ", True)) # Species Italic
            preview_str += f"*{species}* "
            
        if author:
            rich_content.append((f"{author}\n", False))
            preview_str += f"{author}\n"
        else:
            rich_content.append(("\n", False))
            preview_str += "\n"
            
        rich_content.append((f"det. {det_name} {det_year}", False))
        preview_str += f"det. {det_name} {det_year}"
        
        st.session_state.label_queue.append({
            'type': 'rich',
            'content': rich_content,
            'quantity': quantity,
            'preview': f"[ID] {genus} {species}"
        })
        st.success(f"Added {quantity} ID Label(s) to Queue!")
        st.text("Preview Format:")
        st.markdown(preview_str)

# --- TAB 3: MOLECULAR LABEL ---
with tab3:
    st.header("Molecular Label")
    mol_id = st.text_input("Sample ID (e.g. DNA-001)")
    mol_note = st.text_input("Note / Method", value="DNA extracted")
    
    add_mol_btn = st.button("Add Molecular Label", use_container_width=True)
    
    if add_mol_btn:
        if not mol_id:
            st.error("Sample ID is required.")
        else:
            text = f"{mol_id}\n{mol_note}"
            st.session_state.label_queue.append({
                'type': 'text',
                'content': text,
                'quantity': quantity,
                'preview': f"[DNA] {mol_id}"
            })
            st.success(f"Added {quantity} Molecular Label(s) to Queue!")

# --- TAB 4: SHEET PREVIEW (Full A4) ---
with tab4:
    st.header("📄 True Sheet Preview (A4)")
    st.info("This preview simulates the A4 layout. Dotted lines representing cut marks are shown in gray.")
    
    if st.session_state.label_queue:
        if st.button("Refresh Preview"):
            st.rerun()
            
        html_content = generate_html_sheet(
            st.session_state.label_queue,
            num_columns=num_columns,
            font_name=font_name,
            font_size=font_size,
            label_color=label_color
        )
        # Display Scrollable HTML
        st.components.v1.html(html_content, height=800, scrolling=True)
    else:
        st.warning("Queue is empty. Add labels to see the sheet preview.")

# Queue Display & Download (Bottom Section)
st.divider()
st.subheader("📦 Current Batch Queue")

if st.session_state.label_queue:
    queue = st.session_state.label_queue
    total_items = len(queue)
    total_labels = sum(item['quantity'] for item in queue)

    # --- Summary Bar ---
    st.markdown(f"**{total_items}** アイテム / **{total_labels}** ラベル（合計）")

    # --- Slider Navigation ---
    if total_items == 1:
        selected_idx = 0
        st.markdown("**Item 1 / 1**")
    else:
        selected_idx = st.slider(
            "アイテムを選択",
            min_value=1,
            max_value=total_items,
            value=1,
            format="Item %d",
            key="queue_slider"
        ) - 1

    item = queue[selected_idx]
    item_type = item.get('type', 'text')

    # --- Card Display ---
    with st.container():
        # Type Badge
        type_labels = {
            'data_v2': '🌎 Data Label',
            'rich': '🔍 Identification Label',
            'text': '🧬 Molecular Label',
        }
        badge = type_labels.get(item_type, '📄 Label')

        st.markdown(f"### {badge}　`#{selected_idx + 1}`")

        if item_type == 'data_v2':
            # --- Data Label Card ---
            card_col1, card_col2 = st.columns([1.2, 1])

            with card_col1:
                # Parse body into structured fields
                header = item.get('header', '')
                body = item.get('body', '')
                body_lines = body.split('\n')

                # Extract fields from body text
                locality_line = body_lines[0] if len(body_lines) > 0 else ''
                coords_date_line = body_lines[1] if len(body_lines) > 1 else ''
                collector_line = body_lines[2] if len(body_lines) > 2 else ''

                # Parse locality and elevation
                locality_part = locality_line.rstrip(',')
                elev_match = re.search(r'\(alt\.\s*(.+?)\s*m?\)', locality_part)
                if elev_match:
                    elev_str = elev_match.group(1)
                    loc_str = locality_part[:locality_part.index('(alt.')].strip().rstrip(',')
                else:
                    elev_str = '—'
                    loc_str = locality_part

                # Parse coords and date from line 2
                coord_match = re.search(r'([\d.]+°[NS]),\s*([\d.]+°[EW])', coords_date_line)
                date_match = re.search(r'(\d+\s+[IVXLCDM]+\s+\d{4})', coords_date_line)
                coord_str = f"{coord_match.group(1)}, {coord_match.group(2)}" if coord_match else '—'
                date_str = date_match.group(1) if date_match else '—'

                # Parse collector and method from line 3
                method_match = re.search(r',\s*\((.+?)\)\s*$', collector_line)
                if method_match:
                    method_str = method_match.group(1)
                    collector_str = collector_line[:collector_line.rindex(',')].strip()
                else:
                    method_str = '—'
                    collector_str = collector_line.strip()

                color_hex = item.get('color', '#000000')

                # Structured info display
                st.markdown(f"""
| | |
|:---|:---|
| **🏷️ ヘッダー** | {header} |
| **📍 場所** | {loc_str} |
| **⛰️ 標高** | {elev_str} m |
| **🌐 座標** | {coord_str} |
| **📅 日付** | {date_str} |
| **👤 採集者** | {collector_str} |
| **🪤 採集方法** | {method_str} |
| **🎨 カラー** | `{color_hex}` |
""")

            with card_col2:
                # HTML Preview
                st.markdown("**ラベルプレビュー:**")
                preview_html = f"""
                <div style="
                    border: 1px solid #555; padding: 8px; max-width: 280px;
                    font-family: Arial; font-size: 11px; line-height: 1.2;
                    background: white; color: black; border-radius: 4px;
                    box-shadow: 0 2px 8px rgba(0,0,0,0.15);
                ">
                    <div style="font-weight: bold;">{header}</div>
                    <div style="height: 4px; background-color: {color_hex}; margin: 3px 0; border-radius: 2px;"></div>
                    <div style="white-space: pre-wrap;">{body}</div>
                </div>
                """
                st.components.v1.html(preview_html, height=140)

        elif item_type == 'rich':
            # --- ID Label Card ---
            content_parts = item.get('content', [])
            family_str = ''
            genus_str = ''
            species_str = ''
            author_str = ''
            det_str = ''

            for text, is_italic in content_parts:
                text_clean = text.strip()
                if not text_clean:
                    continue
                if text_clean.startswith('det.'):
                    det_str = text_clean
                elif is_italic and not genus_str:
                    genus_str = text_clean
                elif is_italic:
                    species_str = text_clean
                elif not family_str and not genus_str:
                    family_str = text_clean
                else:
                    author_str = text_clean

            st.markdown(f"""
| | |
|:---|:---|
| **Family** | {family_str or '—'} |
| **Genus** | *{genus_str}* |
| **Species** | *{species_str or '—'}* |
| **Author** | {author_str or '—'} |
| **Det.** | {det_str or '—'} |
""")

            # Preview
            preview_parts = []
            for text, is_italic in content_parts:
                if is_italic:
                    preview_parts.append(f"*{text.strip()}*")
                else:
                    preview_parts.append(text.strip())
            st.markdown("**プレビュー:** " + " ".join([p for p in preview_parts if p]))

        else:
            # --- Molecular / Text Label Card ---
            content = item.get('content', '')
            lines = str(content).split('\n')
            sample_id = lines[0] if len(lines) > 0 else '—'
            note = lines[1] if len(lines) > 1 else '—'

            st.markdown(f"""
| | |
|:---|:---|
| **🧪 Sample ID** | `{sample_id}` |
| **📝 Note** | {note} |
""")

        # --- Action Buttons ---
        st.markdown("---")
        act_col1, act_col2, act_col3, act_col4, act_col5 = st.columns([1, 0.5, 0.5, 0.5, 1])

        with act_col1:
            st.markdown(f"**数量: {item['quantity']}**")

        with act_col2:
            if st.button("➖", key=f"qty_minus_{selected_idx}", help="数量を減らす"):
                if item['quantity'] > 1:
                    queue[selected_idx]['quantity'] -= 1
                    st.rerun()

        with act_col3:
            if st.button("➕", key=f"qty_plus_{selected_idx}", help="数量を増やす"):
                queue[selected_idx]['quantity'] += 1
                st.rerun()

        with act_col5:
            if st.button("🗑️ 削除", key=f"del_{selected_idx}", type="secondary"):
                queue.pop(selected_idx)
                st.rerun()

    st.divider()

    # --- Download Batch ---
    docx_file = create_docx(
        queue,
        font_size=font_size,
        show_borders=show_borders,
        num_columns=num_columns,
        font_name=font_name,
        char_spacing=char_spacing
    )
    st.download_button(
        label=f"📥 Download Batch DOCX ({total_items} types / {total_labels} labels)",
        data=docx_file,
        file_name=f"labels_batch_{datetime.date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )

    # --- Summary Table (Collapsible) ---
    with st.expander("📋 全アイテム一覧", expanded=False):
        summary_data = []
        for i, item in enumerate(queue):
            item_type = item.get('type', 'text')
            type_name = {'data_v2': 'Data', 'rich': 'ID', 'text': 'Molecular'}.get(item_type, 'Other')
            preview = item.get('preview', '').replace('\n', ' ')
            if len(preview) > 60:
                preview = preview[:60] + '...'
            summary_data.append({
                '#': i + 1,
                'Type': type_name,
                'Preview': preview,
                'Qty': item['quantity'],
            })
        df = pd.DataFrame(summary_data)
        st.dataframe(df, use_container_width=True, hide_index=True)

else:
    st.info("Queue is empty. Add labels from the tabs above.")
